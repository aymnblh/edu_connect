from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("D:/Aymen/edu/artifacts/Presentation_EduConnect_Algerie.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 36, 61)
MUTED = RGBColor(95, 110, 130)
TEAL = RGBColor(15, 118, 110)
LIGHT_FILL = "F4F6F9"
SOFT_TEAL = "E6F4F1"
SOFT_BLUE = "EAF2FB"
TABLE_BORDER = "D7DBE2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=TABLE_BORDER, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(paragraph, *, before=0, after=8, line=1.333, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_para(doc, text="", *, size=11, color=INK, bold=False, italic=False, after=8, before=0, align=None):
    p = doc.add_paragraph()
    style_paragraph(p, before=before, after=after, align=align)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        style_paragraph(p, before=18, after=10, line=1.2)
        size, color = 16, BLUE
    elif level == 2:
        style_paragraph(p, before=12, after=6, line=1.2)
        size, color = 13, BLUE
    else:
        style_paragraph(p, before=8, after=4, line=1.2)
        size, color = 12, DARK_BLUE
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, before=0, after=4, line=1.208)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, before=0, after=4, line=1.208)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_callout(doc, label, text, fill=SOFT_TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.25)
    r1 = p.add_run(label)
    set_run_font(r1, size=11, color=TEAL, bold=True)
    r2 = p.add_run(f" {text}")
    set_run_font(r2, size=11, color=INK)
    add_para(doc, "", after=4)


def apply_base_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.right_margin = Inches(0.95)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.95)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = INK
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def set_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("EduConnect | Présentation et bénéfices pour les établissements")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Document de présentation - usage commercial et administratif")
    set_run_font(r, size=8.5, color=MUTED)


def add_cover(doc):
    add_para(doc, "EduConnect", size=13, color=TEAL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    p = add_para(
        doc,
        "Une application scolaire moderne pour rapprocher l'établissement, les enseignants et les familles",
        size=24,
        color=INK,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
    )
    p.paragraph_format.line_spacing = 1.15
    add_para(
        doc,
        "Présentation simple pour les responsables d'établissement, les parents et les équipes pédagogiques en Algérie",
        size=13,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=20,
    )
    add_callout(
        doc,
        "Idée principale :",
        "EduConnect donne à l'école une image plus organisée, plus proche des parents et plus moderne, tout en simplifiant le suivi quotidien des élèves.",
    )
    add_para(
        doc,
        "Ce document explique, sans langage technique, comment l'application s'utilise et pourquoi elle peut devenir un vrai avantage pour un établissement privé ou public qui souhaite améliorer sa communication.",
        size=11,
        color=INK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=14,
    )


def add_benefit_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [2450, 3455, 3455])
    set_cell_margins(table, top=120, bottom=120, start=140, end=140)
    headers = ["Pour qui", "Ce que l'application simplifie", "Bénéfice concret"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.15)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)

    rows = [
        (
            "Direction",
            "Suivi global de l'école, des classes, des enseignants, des familles, des paiements et de l'activité.",
            "Moins de dispersion, meilleure visibilité et décisions plus rapides.",
        ),
        (
            "Enseignants",
            "Appel, notes, devoirs, remarques et messages aux familles.",
            "Gain de temps et moins d'allers-retours inutiles.",
        ),
        (
            "Parents",
            "Consultation des notes, absences, devoirs, messages et informations importantes.",
            "Plus besoin d'attendre ou de se déplacer pour chaque information.",
        ),
        (
            "Secrétariat",
            "Import des élèves, génération des codes parents, organisation administrative.",
            "Moins de papier, moins d'erreurs, meilleur suivi.",
        ),
        (
            "Élèves",
            "Suivi plus clair entre l'école et la maison.",
            "Encadrement plus régulier et réaction plus rapide en cas de difficulté.",
        ),
    ]
    for who, simple, benefit in rows:
        cells = table.add_row().cells
        values = [who, simple, benefit]
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            style_paragraph(p, after=0, line=1.2)
            r = p.add_run(value)
            set_run_font(r, size=10.2, color=INK, bold=(idx == 0))
    add_para(doc, "", after=4)


def add_examples_table(doc):
    examples = [
        ("Absence d'un élève", "Le parent est informé plus vite et l'administration garde une trace claire."),
        ("Nouvelle note", "Le parent voit l'évolution sans attendre le bulletin papier."),
        ("Devoir important", "L'enseignant publie le devoir et les familles peuvent le consulter à tout moment."),
        ("Message à une famille", "Le contact passe par un canal scolaire officiel, pas par un groupe non contrôlé."),
        ("Gestion des paiements", "L'établissement peut suivre les factures et paiements plus clairement."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [3000, 6360])
    set_cell_margins(table, top=120, bottom=120, start=140, end=140)
    for idx, header in enumerate(["Situation", "Ce qu'EduConnect apporte"]):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.15)
        r = p.add_run(header)
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    for situation, contribution in examples:
        cells = table.add_row().cells
        for idx, value in enumerate([situation, contribution]):
            p = cells[idx].paragraphs[0]
            style_paragraph(p, after=0, line=1.2)
            r = p.add_run(value)
            set_run_font(r, size=10.2, color=INK, bold=(idx == 0))
    add_para(doc, "", after=4)


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    apply_base_styles(doc)
    set_header_footer(doc)
    add_cover(doc)

    add_heading(doc, "1. Pourquoi une application scolaire aujourd'hui ?", 1)
    add_para(
        doc,
        "Dans beaucoup d'établissements en Algérie, une grande partie de la communication repose encore sur les appels téléphoniques, les carnets, les groupes informels et les déplacements des parents. Ces moyens restent utiles, mais ils deviennent vite difficiles à suivre lorsque le nombre d'élèves augmente.",
    )
    add_para(
        doc,
        "EduConnect centralise les informations importantes dans un seul espace. L'établissement garde le contrôle, les enseignants communiquent plus simplement, et les parents suivent la scolarité de leurs enfants sans attendre la fin du trimestre.",
    )
    add_callout(
        doc,
        "Promesse :",
        "moins d'oublis, moins de confusion, plus de confiance entre l'école et les familles.",
    )

    add_heading(doc, "2. Comment l'application s'utilise au quotidien", 1)
    add_para(doc, "L'utilisation est pensée pour rester simple. Chaque personne voit uniquement ce qui la concerne.")
    for item in [
        "La direction crée les classes, ajoute les élèves et invite les enseignants.",
        "L'établissement remet au parent un QR code, un identifiant élève ou un code PIN pour relier son enfant en toute sécurité.",
        "L'enseignant fait l'appel, publie les notes, ajoute les devoirs et peut envoyer un message aux familles.",
        "Le parent ouvre l'application pour consulter les absences, les notes, les devoirs et les messages importants.",
        "Les informations restent organisées par école, par classe et par rôle.",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "3. Ce qu'EduConnect change pour chaque personne", 1)
    add_benefit_table(doc)

    add_heading(doc, "4. Bénéfices pour les parents", 1)
    add_para(
        doc,
        "Pour un parent, le plus grand avantage est la tranquillité. Il n'a plus besoin d'attendre une réunion ou un appel pour savoir si son enfant est présent, s'il a reçu une note, ou s'il y a un devoir important.",
    )
    for item in [
        "suivre les notes et absences plus rapidement ;",
        "recevoir les informations importantes de l'école dans un espace clair ;",
        "communiquer avec l'équipe pédagogique sans passer par des canaux informels ;",
        "mieux accompagner l'enfant à la maison ;",
        "réduire les déplacements inutiles vers l'établissement.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. Bénéfices pour l'établissement", 1)
    add_para(
        doc,
        "Pour l'établissement, EduConnect n'est pas seulement une application. C'est un outil d'organisation et d'image. Il montre que l'école prend au sérieux la communication, la sécurité des données scolaires et la relation avec les familles.",
    )
    for item in [
        "une administration plus fluide et moins dépendante du papier ;",
        "une meilleure relation avec les parents, surtout dans les moments sensibles ;",
        "un suivi plus rapide des absences, remarques et difficultés scolaires ;",
        "une valorisation de l'établissement face aux familles qui comparent les écoles ;",
        "une image d'école moderne, structurée et attentive.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Pourquoi c'est particulièrement utile en Algérie", 1)
    add_para(
        doc,
        "En Algérie, les familles accordent une grande importance au suivi scolaire et à la réputation de l'établissement. Une application comme EduConnect répond à un besoin très concret : garder un lien régulier entre l'école et la maison, sans compliquer la vie des parents.",
    )
    add_para(
        doc,
        "Elle aide aussi les écoles à se différencier. Pour un parent qui cherche un établissement sérieux, la présence d'un outil numérique bien organisé donne un signal fort : l'école est professionnelle, transparente et proche des familles.",
    )
    for item in [
        "les parents sont souvent très impliqués dans le parcours scolaire ;",
        "les déplacements peuvent prendre du temps, surtout dans les grandes villes ;",
        "les groupes de messagerie informels peuvent créer de la confusion ;",
        "les établissements ont besoin de se distinguer par la qualité du suivi ;",
        "la confiance se construit par une information rapide, claire et vérifiable.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. L'image que l'application donne à l'établissement", 1)
    add_callout(
        doc,
        "Image perçue :",
        "une école équipée d'EduConnect paraît plus moderne, plus transparente, plus organisée et plus proche des familles.",
        fill=SOFT_BLUE,
    )
    add_para(
        doc,
        "Cette image compte beaucoup. Elle rassure les parents avant même qu'ils utilisent toutes les fonctionnalités. Elle montre que l'établissement ne se contente pas d'enseigner, mais qu'il accompagne aussi la famille dans le suivi de l'enfant.",
    )
    for item in [
        "Sérieux : les informations scolaires sont structurées et suivies.",
        "Modernité : l'école adopte des outils actuels et utiles.",
        "Proximité : les parents sentent que l'établissement reste accessible.",
        "Confiance : les données et les messages passent par un canal officiel.",
        "Professionnalisme : l'école maîtrise son organisation interne.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Exemples simples de situations où l'application facilite la vie", 1)
    add_examples_table(doc)

    add_heading(doc, "9. Message de conclusion pour présenter l'application", 1)
    add_para(
        doc,
        "EduConnect aide l'établissement à gagner du temps, à mieux informer les familles et à donner une image plus professionnelle. Pour les parents, c'est un moyen simple de rester proches de la scolarité de leurs enfants. Pour les enseignants, c'est un outil qui facilite les gestes du quotidien. Pour la direction, c'est une preuve d'organisation et de modernité.",
    )
    add_callout(
        doc,
        "En une phrase :",
        "EduConnect transforme la communication scolaire en un service clair, moderne et rassurant pour toute la communauté éducative.",
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
